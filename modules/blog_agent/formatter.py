# modules/blog_agent/formatter.py
import os
import re
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from modules.utils import ensure_dir

# --- Define Project Root to find assets/logo.jpg ---
# This makes the path robust, finding D:\Marketing Agent\assets\logo.jpg
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
LOGO_PATH = os.path.join(PROJECT_ROOT, "assests", "logo.jpg")


def add_hyperlink(paragraph, text, url):
    """
    A function to add a clickable hyperlink to a paragraph.
    This is the standard, correct method using OXML.
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Apply the 'Hyperlink' style (blue, underlined)
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    
    # Set font size to 10pt (value is in half-points, so 20)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')
    rPr.append(sz)
    
    # Apply Italic
    i = OxmlElement('w:i')
    rPr.append(i)

    new_run.append(rPr)
    
    # --- START OF FIX ---
    # We create the text element separately and append it
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    # --- END OF FIX ---
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def handle_formatting(paragraph, text):
    """
    Parses a line of text for **bold** and *italic* Markdown.
    """
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            sub_parts = re.split(r'(\*.*?\*)', content)
            for sub_part in sub_parts:
                if sub_part.startswith('*') and sub_part.endswith('*'):
                    paragraph.add_run(sub_part[1:-1]).italic = True
                else:
                    paragraph.add_run(sub_part).bold = True
        else:
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sub_part in sub_parts:
                if sub_part.startswith('*') and sub_part.endswith('*'):
                    paragraph.add_run(sub_part[1:-1]).italic = True
                else:
                    paragraph.add_run(sub_part)

def add_branding_footer(document):
    """
    Adds a branded footer with logo and clickable hyperlink.
    """
    BRAND_TEXT = "@aiwithsid | "
    BRAND_URL = "http://grwothbrothers.in"
    
    section = document.sections[0]
    footer = section.footer
    
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = "" 
    
    if os.path.exists(LOGO_PATH):
        try:
            run_logo = p.add_run()
            run_logo.add_picture(LOGO_PATH, height=Inches(0.3))
        except Exception as e:
            print(f"⚠️  Could not add logo to DOCX footer: {e}")
    else:
        # === PATH FIX IS HERE ===
        print(f"⚠️  Logo not found at {LOGO_PATH}. Skipping footer logo.")
        
    # === FONT SIZE FIX IS HERE ===
    run_text = p.add_run(f"\t{BRAND_TEXT}")
    run_text.font.size = Pt(10) # Increased from 9 to 10
    run_text.italic = True
    
    # === HYPERLINK FIX IS HERE ===
    # Add the clickable hyperlink
    add_hyperlink(p, BRAND_URL, BRAND_URL)
    
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_default_font(document):
    """
    Sets the default font for the document.
    """
    try:
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)
    except Exception as e:
        print(f"⚠️  Could not set default font. Using document defaults. Error: {e}")

def assemble_docx(plan: dict, sections_with_md: list[tuple[str,str]], cover_path: str|None, topic: str, run_id: str = None) -> str:
    """
    Assemble the final blog post as a .docx file with a UNIQUE filename.
    """
    title = plan.get("title", f"Understanding {topic}")
    out_dir = "generated/blogs"
    assets_dir = os.path.join(out_dir, "assets")
    ensure_dir(os.path.join(assets_dir, "x"))
    
    # Generate unique filename
    if run_id:
        docx_filename = f"blog_{run_id}.docx"
    else:
        docx_filename = "blog.docx"
        
    docx_path = os.path.join(out_dir, docx_filename)

    document = Document()
    set_default_font(document)

    # --- Header ---
    document.add_heading(title, level=1)
    p = document.add_paragraph()
    p.add_run(f"Published: {date.today().isoformat()}").italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- Cover Image ---
    if cover_path and os.path.exists(cover_path):
        try:
            document.add_picture(cover_path, width=Inches(6.0))
            document.add_paragraph()
        except Exception as e:
            print(f"⚠️  Could not add cover image to DOCX: {e}")

    # --- Blog Sections ---
    img_regex = re.compile(r'!\[(.*?)\]\((.*?)\)')

    for heading, content_md in sections_with_md:
        if "introduction" not in heading.lower():
            document.add_heading(heading, level=2)

        for line in content_md.split('\n'):
            line = line.strip()
            if not line:
                continue

            match = img_regex.search(line)
            if match:
                img_desc, img_rel_path = match.groups()
                # Fix path resolution for relative paths
                img_full_path = os.path.join(out_dir, img_rel_path.replace("/", os.sep))
                
                if os.path.exists(img_full_path):
                    try:
                        document.add_picture(img_full_path, width=Inches(5.5))
                        if img_desc:
                             p = document.add_paragraph()
                             p.add_run(f"Figure: {img_desc}").italic = True
                             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"⚠️  Could not add section image '{img_full_path}': {e}")
                continue

            if line.startswith('## '):
                document.add_heading(line.lstrip('## ').strip(), level=3)
            elif line.startswith('* '):
                p = document.add_paragraph(style='List Bullet')
                handle_formatting(p, line.lstrip('* ').strip())
            elif re.match(r'^\d+\.\s', line):
                 p = document.add_paragraph(style='List Number')
                 handle_formatting(p, re.sub(r'^\d+\.\s', '', line).strip())
            else:
                p = document.add_paragraph()
                handle_formatting(p, line)

    # --- Add Branding Footer ---
    add_branding_footer(document)

    # --- Save the document ---
    try:
        document.save(docx_path)
        print(f"✅ Successfully created DOCX file: {docx_path}")
    except Exception as e:
        print(f"❌ Failed to save DOCX file: {e}")

    return docx_path